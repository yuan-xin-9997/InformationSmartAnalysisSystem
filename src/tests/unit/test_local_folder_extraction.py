"""TDD tests for local_folder metadata + figure extraction (task 2).

Covers spec scenarios:
- PDF metadata extraction (title/author/creationDate/page_count)
- Author-affiliation line matching (hit / no-hit / empty / English)
- Figure extraction (PDF multi-image / no-image / truncate / docx / txt empty)
- _extract_full packs metadata + figures into InfoItemData.extra
- fetch_new_items carries metadata
- reextract returns fresh data or None for missing files
- txt/md fallback to filename, all metadata None
"""
from __future__ import annotations

import struct
import zlib
from pathlib import Path

import pytest


# ---------- helpers to build sample documents ----------


def _make_png(width: int = 2, height: int = 3, color: tuple = (255, 0, 0)) -> bytes:
    """Create a minimal valid PNG (RGB) with known dimensions."""
    sig = b"\x89PNG\r\n\x1a\n"
    # IHDR: width, height, bit_depth=8, color_type=2 (RGB), compression=0, filter=0, interlace=0
    ihdr_data = struct.pack(">IIBBBBB", width, height, 8, 2, 0, 0, 0)
    ihdr = b"IHDR" + ihdr_data
    ihdr_chunk = (
        struct.pack(">I", len(ihdr_data))
        + ihdr
        + struct.pack(">I", zlib.crc32(ihdr) & 0xFFFFFFFF)
    )
    # IDAT: one filter byte (0) per scanline + RGB pixels
    raw = b""
    px = bytes(color)
    for _ in range(height):
        raw += b"\x00" + px * width
    compressed = zlib.compress(raw)
    idat = b"IDAT" + compressed
    idat_chunk = (
        struct.pack(">I", len(compressed))
        + idat
        + struct.pack(">I", zlib.crc32(idat) & 0xFFFFFFFF)
    )
    # IEND
    iend = b"IEND"
    iend_chunk = (
        struct.pack(">I", 0)
        + iend
        + struct.pack(">I", zlib.crc32(iend) & 0xFFFFFFFF)
    )
    return sig + ihdr_chunk + idat_chunk + iend_chunk


def _make_pdf(
    path: Path,
    title: str | None = "T",
    author: str | None = "A",
    creation_date: str | None = "D:20260115103000+00'00'",
    text_lines: list[str] | None = None,
    image_paths: list[Path] | None = None,
) -> None:
    import fitz

    doc = fitz.open()
    md: dict[str, str] = {}
    if title is not None:
        md["title"] = title
    if author is not None:
        md["author"] = author
    if creation_date is not None:
        md["creationDate"] = creation_date
    if md:
        doc.set_metadata(md)
    page = doc.new_page()
    if text_lines:
        page.insert_text((50, 50), "\n".join(text_lines), fontname="china-s")
    if image_paths:
        for img in image_paths:
            rect = fitz.Rect(0, 0, 100, 100)
            page.insert_image(rect, filename=str(img))
    doc.save(str(path))
    doc.close()


def _make_docx(
    path: Path,
    title: str | None = "T",
    author: str | None = "A",
    text: str = "body",
    image_paths: list[Path] | None = None,
) -> None:
    from docx import Document
    from docx.shared import Inches

    doc = Document()
    if text:
        doc.add_paragraph(text)
    if image_paths:
        for img in image_paths:
            doc.add_picture(str(img), width=Inches(1))
    if title is not None:
        doc.core_properties.title = title
    if author is not None:
        doc.core_properties.author = author
    doc.save(str(path))


def _make_html(
    path: Path,
    title: str = "T",
    author: str = "A",
    published: str = "2026-01-15T10:30:00+00:00",
    body: str = "<p>hello</p>",
) -> None:
    html = (
        f"<!DOCTYPE html>\n<html><head><title>{title}</title>\n"
        f'<meta name="author" content="{author}">\n'
        f'<meta property="article:published_time" content="{published}">\n'
        f"</head><body>{body}</body></html>"
    )
    path.write_text(html, encoding="utf-8")


# ---------- extract_metadata ----------


def test_extract_metadata_pdf(tmp_path):
    from app.backend.services.info_source.local_folder import extract_metadata

    pdf = tmp_path / "a.pdf"
    _make_pdf(pdf, title="My Title", author="John", creation_date="D:20260115103000+00'00'")
    md = extract_metadata(pdf)
    assert md["title"] == "My Title"
    assert md["author"] == "John"
    assert md["page_count"] == 1
    assert md["published_at"] is not None
    assert md["published_at"].year == 2026
    assert md["published_at"].month == 1
    assert md["published_at"].day == 15


def test_extract_metadata_pdf_no_metadata(tmp_path):
    from app.backend.services.info_source.local_folder import extract_metadata

    pdf = tmp_path / "a.pdf"
    _make_pdf(pdf, title=None, author=None, creation_date=None)
    md = extract_metadata(pdf)
    assert md["title"] is None
    assert md["author"] is None
    assert md["page_count"] == 1  # page_count is always derivable for PDF
    assert md["published_at"] is None


def test_extract_metadata_docx(tmp_path):
    from app.backend.services.info_source.local_folder import extract_metadata

    docx_path = tmp_path / "a.docx"
    _make_docx(docx_path, title="Doc Title", author="Jane")
    md = extract_metadata(docx_path)
    assert md["title"] == "Doc Title"
    assert md["author"] == "Jane"
    assert md["page_count"] is None  # docx has no native page count
    # core_properties.created is set by python-docx automatically
    assert md["published_at"] is not None


def test_extract_metadata_html(tmp_path):
    from app.backend.services.info_source.local_folder import extract_metadata

    html = tmp_path / "a.html"
    _make_html(html, title="HTML Title", author="Web Author")
    md = extract_metadata(html)
    assert md["title"] == "HTML Title"
    assert md["author"] == "Web Author"
    assert md["page_count"] is None
    assert md["published_at"] is not None
    assert md["published_at"].year == 2026


def test_extract_metadata_txt_all_none(tmp_path):
    from app.backend.services.info_source.local_folder import extract_metadata

    txt = tmp_path / "a.txt"
    txt.write_text("hello", encoding="utf-8")
    md = extract_metadata(txt)
    assert md["title"] is None
    assert md["author"] is None
    assert md["published_at"] is None
    assert md["page_count"] is None


def test_extract_metadata_md_all_none(tmp_path):
    from app.backend.services.info_source.local_folder import extract_metadata

    md_file = tmp_path / "a.md"
    md_file.write_text("# title\nbody", encoding="utf-8")
    md = extract_metadata(md_file)
    assert md["title"] is None
    assert md["author"] is None
    assert md["published_at"] is None
    assert md["page_count"] is None


# ---------- extract_author_affiliation ----------


def test_affiliation_hit_chinese():
    from app.backend.services.info_source.local_folder import extract_author_affiliation

    text = "Title\n北京大学 计算机学院\nSome other text"
    result = extract_author_affiliation(text)
    assert result is not None
    assert "北京大学" in result


def test_affiliation_hit_english():
    from app.backend.services.info_source.local_folder import extract_author_affiliation

    text = "Title\nDepartment of CS, Some University\nbody"
    result = extract_author_affiliation(text)
    assert result is not None
    assert "University" in result


def test_affiliation_no_hit():
    from app.backend.services.info_source.local_folder import extract_author_affiliation

    assert extract_author_affiliation("Just some text\nNo keywords here") is None


def test_affiliation_empty_and_none():
    from app.backend.services.info_source.local_folder import extract_author_affiliation

    assert extract_author_affiliation("") is None
    assert extract_author_affiliation(None) is None  # type: ignore[arg-type]


def test_affiliation_first_hit_wins():
    from app.backend.services.info_source.local_folder import extract_author_affiliation

    text = "Title\n清华大学 计算机系\n北京大学 物理系"
    result = extract_author_affiliation(text)
    assert result is not None
    assert "清华大学" in result
    assert "北京大学" not in result


# ---------- extract_figures ----------


def test_extract_figures_pdf_multiple(tmp_path):
    from app.backend.services.info_source.local_folder import extract_figures

    img1 = tmp_path / "i1.png"
    img1.write_bytes(_make_png(2, 3))
    img2 = tmp_path / "i2.png"
    img2.write_bytes(_make_png(4, 5))
    img3 = tmp_path / "i3.png"
    img3.write_bytes(_make_png(6, 7))
    pdf = tmp_path / "a.pdf"
    _make_pdf(pdf, image_paths=[img1, img2, img3])
    figs = extract_figures(pdf, max_count=20)
    assert len(figs) == 3
    assert all(f.bytes_data for f in figs)
    # width/height parsed from PNG header (order may vary by xref)
    dims = sorted((f.width, f.height) for f in figs)
    assert (2, 3) in dims
    assert (4, 5) in dims
    assert (6, 7) in dims
    assert all(f.mime == "image/png" for f in figs)


def test_extract_figures_pdf_none(tmp_path):
    from app.backend.services.info_source.local_folder import extract_figures

    pdf = tmp_path / "a.pdf"
    _make_pdf(pdf)
    figs = extract_figures(pdf, max_count=20)
    assert figs == []


def test_extract_figures_pdf_truncate(tmp_path):
    from app.backend.services.info_source.local_folder import extract_figures

    img1 = tmp_path / "i1.png"
    img1.write_bytes(_make_png(2, 3))
    img2 = tmp_path / "i2.png"
    img2.write_bytes(_make_png(4, 5))
    img3 = tmp_path / "i3.png"
    img3.write_bytes(_make_png(6, 7))
    pdf = tmp_path / "a.pdf"
    _make_pdf(pdf, image_paths=[img1, img2, img3])
    figs = extract_figures(pdf, max_count=2)
    assert len(figs) == 2


def test_extract_figures_docx(tmp_path):
    from app.backend.services.info_source.local_folder import extract_figures

    img = tmp_path / "img.png"
    img.write_bytes(_make_png(4, 5))
    docx_path = tmp_path / "a.docx"
    _make_docx(docx_path, image_paths=[img])
    figs = extract_figures(docx_path, max_count=20)
    assert len(figs) == 1
    assert figs[0].width == 4
    assert figs[0].height == 5
    assert figs[0].mime == "image/png"


def test_extract_figures_txt_empty(tmp_path):
    from app.backend.services.info_source.local_folder import extract_figures

    txt = tmp_path / "a.txt"
    txt.write_text("hello", encoding="utf-8")
    assert extract_figures(txt, max_count=20) == []


def test_extract_figures_html_empty(tmp_path):
    from app.backend.services.info_source.local_folder import extract_figures

    html = tmp_path / "a.html"
    html.write_text("<html><body><p>hi</p></body></html>", encoding="utf-8")
    assert extract_figures(html, max_count=20) == []


# ---------- _extract_full ----------


def test_extract_full_pdf_with_metadata_and_figures(tmp_path):
    from app.backend.services.info_source.local_folder import _extract_full

    img = tmp_path / "img.png"
    img.write_bytes(_make_png(2, 3))
    pdf = tmp_path / "doc.pdf"
    _make_pdf(
        pdf,
        title="Title",
        author="Author",
        text_lines=["北京大学 计算机学院"],
        image_paths=[img],
    )
    data = _extract_full(pdf)
    assert data is not None
    assert data.title == "Title"
    assert data.extra["author"] == "Author"
    assert data.extra["author_affiliation"] is not None
    assert "北京大学" in data.extra["author_affiliation"]
    assert data.extra["page_count"] == 1
    assert data.extra["article_published_at"] is not None
    assert len(data.extra["figures"]) == 1
    assert data.extra["figures"][0].width == 2


def test_extract_full_pdf_no_affiliation(tmp_path):
    """PDF with metadata but no affiliation keyword on first page -> affiliation None."""
    from app.backend.services.info_source.local_folder import _extract_full

    pdf = tmp_path / "doc.pdf"
    _make_pdf(pdf, title="T", author="A", text_lines=["just some random text"])
    data = _extract_full(pdf)
    assert data is not None
    assert data.extra["author"] == "A"
    assert data.extra["author_affiliation"] is None
    assert data.extra["page_count"] == 1


def test_extract_full_txt_fallback_filename(tmp_path):
    from app.backend.services.info_source.local_folder import _extract_full

    txt = tmp_path / "notes.txt"
    txt.write_text("just text", encoding="utf-8")
    data = _extract_full(txt)
    assert data is not None
    assert data.title == "notes.txt"  # fallback to filename
    assert data.extra["author"] is None
    assert data.extra["author_affiliation"] is None
    assert data.extra["article_published_at"] is None
    assert data.extra["page_count"] is None
    assert data.extra["figures"] == []


# ---------- fetch_new_items carries metadata ----------


def test_fetch_new_items_carries_metadata(tmp_path):
    from app.backend.services.info_source.local_folder import LocalFolderAdapter

    pdf = tmp_path / "a.pdf"
    _make_pdf(pdf, title="T", author="A", text_lines=["清华大学"])
    adapter = LocalFolderAdapter({"folder_path": str(tmp_path), "patterns": ["*.pdf"]})
    items = adapter.fetch_new_items()
    assert len(items) == 1
    it = items[0]
    assert it.title == "T"
    assert it.extra["author"] == "A"
    assert it.extra["author_affiliation"] is not None
    assert "清华大学" in it.extra["author_affiliation"]
    assert it.extra["page_count"] == 1


def test_fetch_new_items_preserves_incremental_skip(tmp_path):
    """Incremental skip logic must still work after refactor."""
    from datetime import datetime, timezone

    from app.backend.services.info_source.local_folder import LocalFolderAdapter

    (tmp_path / "a.txt").write_text("A", encoding="utf-8")
    adapter = LocalFolderAdapter({"folder_path": str(tmp_path), "patterns": ["*.txt"]})
    first = adapter.fetch_new_items()
    assert len(first) == 1
    known = {it.external_id for it in first}
    since = datetime.now(timezone.utc)
    # Known and unchanged -> skipped
    assert adapter.fetch_new_items(since=since, known_ids=known) == []


# ---------- reextract ----------


def test_reextract_returns_data(tmp_path):
    from app.backend.services.info_source.local_folder import LocalFolderAdapter

    pdf = tmp_path / "a.pdf"
    _make_pdf(pdf, title="T", author="A")
    adapter = LocalFolderAdapter({"folder_path": str(tmp_path), "patterns": ["*.pdf"]})
    data = adapter.reextract(str(pdf.resolve()))
    assert data is not None
    assert data.title == "T"
    assert data.extra["author"] == "A"


def test_reextract_missing_file_returns_none(tmp_path):
    from app.backend.services.info_source.local_folder import LocalFolderAdapter

    adapter = LocalFolderAdapter({"folder_path": str(tmp_path), "patterns": ["*.pdf"]})
    assert adapter.reextract(str(tmp_path / "nope.pdf")) is None


def test_base_adapter_default_reextract_returns_none():
    """Adapters that don't override reextract return None."""
    from app.backend.services.info_source.base import InfoSourceAdapter

    # Can't instantiate ABC directly, but the default method exists.
    # Use a concrete subclass that doesn't override reextract.
    class _Stub(InfoSourceAdapter):
        type = "stub"

        def check_status(self):
            from .base import SourceStatus

            return SourceStatus(ok=True)

        def fetch_new_items(self, since=None, known_ids=None):
            return []

    stub = _Stub({})
    assert stub.reextract("any-id") is None


# ---------- extraction failure logging (review fix) ----------


def test_extract_metadata_logs_warning_on_corrupted_pdf(tmp_path, caplog):
    """A corrupted PDF must emit a warning (not silently return all-None)."""
    import logging

    from app.backend.services.info_source.local_folder import extract_metadata

    pdf = tmp_path / "broken.pdf"
    pdf.write_bytes(b"not a real pdf")
    with caplog.at_level(logging.WARNING, logger="local_folder"):
        md = extract_metadata(pdf)
    assert md == {"title": None, "author": None, "published_at": None, "page_count": None}
    assert any(
        r.levelno == logging.WARNING and "抽取元数据失败" in r.getMessage()
        for r in caplog.records
    )


def test_extract_figures_logs_warning_on_corrupted_pdf(tmp_path, caplog):
    """A corrupted PDF must emit a warning from extract_figures (not silent [])."""
    import logging

    from app.backend.services.info_source.local_folder import extract_figures

    pdf = tmp_path / "broken.pdf"
    pdf.write_bytes(b"not a real pdf")
    with caplog.at_level(logging.WARNING, logger="local_folder"):
        figs = extract_figures(pdf, max_count=10)
    assert figs == []
    assert any(
        r.levelno == logging.WARNING and "抽取图表失败" in r.getMessage()
        for r in caplog.records
    )


def test_first_page_text_logs_warning_on_corrupted_pdf(tmp_path, caplog):
    """A corrupted PDF must emit a warning from _first_page_text (not silent None)."""
    import logging

    from app.backend.services.info_source.local_folder import _first_page_text

    pdf = tmp_path / "broken.pdf"
    pdf.write_bytes(b"not a real pdf")
    with caplog.at_level(logging.WARNING, logger="local_folder"):
        result = _first_page_text(pdf)
    assert result is None
    assert any(
        r.levelno == logging.WARNING and "抽取首页文本失败" in r.getMessage()
        for r in caplog.records
    )
