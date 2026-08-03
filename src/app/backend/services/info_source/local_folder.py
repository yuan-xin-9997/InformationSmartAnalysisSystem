"""Local-folder information-source adapter.

Extracts plain text, article metadata (title/author/affiliation/published_at/
page_count) and embedded figures from local files (txt/md/html/pdf/docx).
``extract_figures`` returns raw bytes only -- ``sync.py`` owns disk persistence
(it has ``settings.figures_dir`` and the ``item_id``).

For PDFs whose text layer is empty or garbled (scanned / broken-font encoding),
a vision-LLM fallback renders pages to images and extracts text via a
multimodal LLM. The quality of the text-layer output is judged by a readable
character ratio; only when it falls below threshold is the fallback engaged.
"""
from __future__ import annotations

import re
import string
from datetime import datetime, timedelta, timezone
from pathlib import Path

from bs4 import BeautifulSoup

from ...core.config import settings
from ...core.logging import get_logger
from .base import FigureData, InfoItemData, InfoSourceAdapter, SourceStatus

_logger = get_logger("local_folder")

# Institutional keywords for author-affiliation line matching.
_AFFILIATION_RE = re.compile(
    r"(大学|学院|研究所|研究院|公司|实验室|医院|中心|科学院|工程院|"
    r"Department|University|Institute|Lab|College|Hospital|Corporation|Inc)"
)

# PDF date: D:YYYYMMDDHHmmSS[+TZ'HH'mm]
_PDF_DATE_RE = re.compile(
    r"D:(\d{4})(\d{2})(\d{2})(\d{2})?(\d{2})?(\d{2})?([+-]\d{2})?'?(\d{2})?'?"
)

# Readable-character set for text-quality scoring: CJK Unified Ideographs,
# ASCII alphanumerics, and common punctuation (ASCII + CJK). Garbled PDF text
# (Private-Use-Area glyphs, mathematical/geometric symbols, replacement chars)
# falls outside this set and drives the readable ratio down.
_CJK_PUNCT = "，。；：！？、（）《》【】「」“”‘’－…·～"
_READABLE_PUNCT = frozenset(string.punctuation + _CJK_PUNCT)


def _is_readable_char(ch: str) -> bool:
    """True for CJK ideographs, ASCII alphanumerics, or common punctuation."""
    code = ord(ch)
    if 0x4E00 <= code <= 0x9FFF:  # CJK Unified Ideographs
        return True
    if ch in _READABLE_PUNCT:
        return True
    if ch.isascii() and ch.isalnum():  # a-z A-Z 0-9
        return True
    return False


def _readable_ratio(text: str) -> float:
    """Readable (non-whitespace) chars / total non-whitespace chars; 0.0 if empty."""
    non_ws = [ch for ch in text if not ch.isspace()]
    if not non_ws:
        return 0.0
    readable = sum(1 for ch in non_ws if _is_readable_char(ch))
    return readable / len(non_ws)


def _text_quality_ok(text: str, min_length: int, readable_ratio_threshold: float) -> bool:
    """True when the text-layer output is usable (non-empty, long enough, readable).

    Returns False for empty / too-short / garbled text -- the signal to engage
    the vision-LLM fallback.
    """
    if not text:
        return False
    non_ws_len = sum(1 for ch in text if not ch.isspace())
    if non_ws_len < min_length:
        return False
    return _readable_ratio(text) >= readable_ratio_threshold


# ---------- text extraction ----------


def extract_text(path: Path) -> str | None:
    """Extract plain text from a file based on its extension."""
    suffix = path.suffix.lower()
    try:
        if suffix in (".txt", ".md"):
            return path.read_text(encoding="utf-8", errors="ignore")
        if suffix in (".html", ".htm"):
            html = path.read_text(encoding="utf-8", errors="ignore")
            return BeautifulSoup(html, "lxml").get_text("\n", strip=True)
        if suffix == ".pdf":
            return _extract_pdf(path)
        if suffix == ".docx":
            return _extract_docx(path)
    except Exception:
        return None
    return None


def _extract_pdf(path: Path) -> str:
    import fitz  # PyMuPDF

    parts: list[str] = []
    with fitz.open(path) as doc:
        for page in doc:
            parts.append(page.get_text())
    return "\n".join(parts).strip()


def _extract_pdf_content(path: Path, llm_client=None) -> tuple[str | None, str]:
    """Extract PDF body text + ``extraction_method``.

    Tries the text layer first; if its quality is poor (empty/garbled) and the
    vision fallback is enabled, renders pages to images and uses a multimodal
    LLM to extract text. Returns ``(content, method)`` where method is
    ``text_layer`` | ``vision_llm`` | ``none``. ``content`` is ``None`` only
    when the text layer itself cannot be read (corrupt file) so the caller
    skips the item; an empty-but-readable text layer yields ``('', ...)``.
    """
    try:
        text_layer = _extract_pdf(path)
    except Exception:
        _logger.warning("PDF 文本层读取失败: %s", path, exc_info=True)
        return None, "none"

    if _text_quality_ok(
        text_layer or "",
        settings.extraction_min_text_length,
        settings.extraction_readable_ratio,
    ):
        return text_layer, "text_layer"

    if not settings.extraction_vision_fallback:
        _logger.info("PDF 文本层质量不佳且视觉兜底未启用，保留原文本: %s", path)
        return text_layer, "none"

    vision_text = _vision_extract_pdf(path, llm_client)
    if vision_text:
        return vision_text, "vision_llm"
    _logger.warning("视觉兜底未产出文本，保留原文本层: %s", path)
    return text_layer, "none"


def _vision_extract_pdf(path: Path, llm_client=None) -> str:
    """Render PDF pages to images and extract text via a multimodal LLM.

    Returns concatenated page text (may be partial), or ``""`` on failure.
    Never raises -- graceful degradation. Respects ``max_ocr_pages`` and
    ``render_dpi``; logs a warning when pages are truncated.
    """
    import fitz  # noqa: F401  (PyMuPDF; used for rendering below)

    from ..analysis.llm_client import LLMClient, LLMError

    max_pages = settings.extraction_max_ocr_pages
    dpi = settings.extraction_render_dpi
    vision_model = settings.extraction_vision_model or None

    if llm_client is None:
        try:
            llm_client = LLMClient(model=vision_model)
        except LLMError as exc:
            _logger.warning("视觉兜底 LLM 未就绪，跳过: %s", exc)
            return ""

    system = (
        "你是一个文档文本提取助手。请提取并原样输出图片中文档的全部正文文本，"
        "保留段落与换行结构，仅输出文本本身，不要解说、不要补充。"
    )
    parts: list[str] = []
    try:
        with fitz.open(path) as doc:
            total = doc.page_count
            pages_to_render = min(total, max_pages)
            if total > max_pages:
                _logger.warning(
                    "PDF 共 %d 页，视觉兜底仅处理前 %d 页（max_ocr_pages）: %s",
                    total,
                    max_pages,
                    path,
                )
            for i in range(pages_to_render):
                pix = doc[i].get_pixmap(dpi=dpi)
                img_bytes = pix.tobytes("png")
                user = (
                    f"请提取这张文档图片（第 {i + 1} 页，共 {total} 页）的全部正文文本。"
                )
                try:
                    page_text = llm_client.chat_with_images(
                        system, user, [img_bytes]
                    )
                except Exception as exc:  # noqa: BLE001 (per-page degrade)
                    _logger.warning(
                        "视觉兜底第 %d 页提取失败，跳过该页: %s", i + 1, exc
                    )
                    continue
                if page_text and page_text.strip():
                    parts.append(page_text.strip())
    except Exception as exc:  # noqa: BLE001 (render-level degrade)
        _logger.warning("视觉兜底渲染失败: %s (%s)", path, exc)
        return ""
    return "\n".join(parts).strip()


def _extract_docx(path: Path) -> str:
    import docx

    doc = docx.Document(str(path))
    return "\n".join(p.text for p in doc.paragraphs).strip()


# ---------- metadata extraction ----------


def extract_metadata(path: Path) -> dict:
    """Return ``{"title", "author", "published_at", "page_count"}``.

    Values are ``None`` when not derivable. For txt/md all are ``None``.
    """
    suffix = path.suffix.lower()
    try:
        if suffix == ".pdf":
            return _metadata_pdf(path)
        if suffix == ".docx":
            return _metadata_docx(path)
        if suffix in (".html", ".htm"):
            return _metadata_html(path)
    except Exception:
        _logger.warning("抽取元数据失败: %s", path, exc_info=True)
        return {"title": None, "author": None, "published_at": None, "page_count": None}
    return {"title": None, "author": None, "published_at": None, "page_count": None}


def _metadata_pdf(path: Path) -> dict:
    import fitz

    with fitz.open(path) as doc:
        md = doc.metadata or {}
        title = (md.get("title") or "").strip() or None
        author = (md.get("author") or "").strip() or None
        page_count = doc.page_count
        published_at = _parse_pdf_date(md.get("creationDate") or "")
    return {
        "title": title,
        "author": author,
        "published_at": published_at,
        "page_count": page_count,
    }


def _metadata_docx(path: Path) -> dict:
    import docx

    doc = docx.Document(str(path))
    cp = doc.core_properties
    title = (cp.title or "").strip() or None
    author = (cp.author or "").strip() or None
    published_at = cp.created
    if isinstance(published_at, datetime):
        if published_at.tzinfo is None:
            published_at = published_at.replace(tzinfo=timezone.utc)
    else:
        published_at = None
    return {
        "title": title,
        "author": author,
        "published_at": published_at,
        "page_count": None,  # docx has no native page-count concept
    }


def _metadata_html(path: Path) -> dict:
    html = path.read_text(encoding="utf-8", errors="ignore")
    soup = BeautifulSoup(html, "lxml")
    title_el = soup.find("title")
    title = title_el.get_text(strip=True) or None if title_el else None
    author = None
    author_meta = soup.find("meta", attrs={"name": "author"})
    if author_meta and author_meta.get("content"):
        author = author_meta["content"].strip() or None
    published_at = None
    pub_meta = soup.find("meta", attrs={"property": "article:published_time"})
    if pub_meta and pub_meta.get("content"):
        published_at = _parse_iso_date(pub_meta["content"])
    return {
        "title": title,
        "author": author,
        "published_at": published_at,
        "page_count": None,
    }


def _parse_pdf_date(s: str) -> datetime | None:
    """Parse a PDF date string like ``D:YYYYMMDDHHmmSS+HH'mm``."""
    if not s:
        return None
    m = _PDF_DATE_RE.match(s)
    if not m:
        return None
    y, mo, d = int(m.group(1)), int(m.group(2)), int(m.group(3))
    h = int(m.group(4) or 0)
    mi = int(m.group(5) or 0)
    sec = int(m.group(6) or 0)
    tz_str = m.group(7)
    if tz_str:
        sign = 1 if tz_str[0] == "+" else -1
        tz_h = int(tz_str[1:3])
        tz_min = int(m.group(8) or 0)
        offset = timedelta(hours=tz_h, minutes=tz_min) * sign
        tz = timezone(offset)
    else:
        tz = timezone.utc
    try:
        return datetime(y, mo, d, h, mi, sec, tzinfo=tz)
    except ValueError:
        return None


def _parse_iso_date(s: str) -> datetime | None:
    """Parse an ISO-8601 date string (handles trailing ``Z``)."""
    if not s:
        return None
    s = s.strip()
    if s.endswith("Z"):
        s = s[:-1] + "+00:00"
    try:
        return datetime.fromisoformat(s)
    except ValueError:
        return None


# ---------- author-affiliation extraction ----------


def extract_author_affiliation(first_page_text: str | None) -> str | None:
    """Scan lines for an institutional keyword; return the first matching line.

    Returns ``None`` when no line matches or the input is empty.
    """
    if not first_page_text:
        return None
    for line in first_page_text.splitlines():
        line = line.strip()
        if not line:
            continue
        if _AFFILIATION_RE.search(line):
            return " ".join(line.split())  # collapse internal whitespace
    return None


def _first_page_text(path: Path) -> str | None:
    """Return a text snippet from the first page for affiliation scanning."""
    suffix = path.suffix.lower()
    try:
        if suffix == ".pdf":
            import fitz

            with fitz.open(path) as doc:
                if len(doc) == 0:
                    return None
                return doc[0].get_text()
        if suffix == ".docx":
            import docx

            d = docx.Document(str(path))
            return "\n".join(p.text for p in d.paragraphs[:20])
        if suffix in (".html", ".htm"):
            html = path.read_text(encoding="utf-8", errors="ignore")
            soup = BeautifulSoup(html, "lxml")
            body = soup.find("body")
            return body.get_text("\n", strip=True) if body else None
    except Exception:
        _logger.warning("抽取首页文本失败: %s", path, exc_info=True)
        return None
    return None


# ---------- figure extraction ----------


def extract_figures(path: Path, max_count: int) -> tuple[list[FigureData], bool]:
    """Extract embedded images as ``FigureData`` (bytes only, no disk writes).

    Returns ``(figures, truncated)``: at most ``max_count`` items, plus a flag
    that is ``True`` when more image candidates existed but were dropped due to
    the cap. For txt/md/html returns ``([], False)``. ``max_count=0`` yields
    ``([], True)`` when the file contains at least one image candidate.
    """
    suffix = path.suffix.lower()
    try:
        if suffix == ".pdf":
            return _figures_pdf(path, max_count)
        if suffix == ".docx":
            return _figures_docx(path, max_count)
    except Exception:
        _logger.warning("抽取图表失败: %s", path, exc_info=True)
        return [], False
    return [], False


def _figures_pdf(path: Path, max_count: int) -> tuple[list[FigureData], bool]:
    import fitz

    figures: list[FigureData] = []
    seen_xrefs: set[int] = set()
    with fitz.open(path) as doc:
        for page in doc:
            for img in page.get_images(full=True):
                # Check the cap BEFORE appending so max_count=0 returns [] and
                # we never overshoot by one.
                if len(figures) >= max_count:
                    return figures, True  # truncated: another candidate remains
                xref = img[0]
                if xref in seen_xrefs:
                    continue
                seen_xrefs.add(xref)
                try:
                    info = doc.extract_image(xref)
                except Exception:
                    continue
                data: bytes | None = info.get("image")
                if not data:
                    continue
                ext = (info.get("ext") or "bin").lower()
                figures.append(_make_figure_data(data, ext))
    return figures, False


def _figures_docx(path: Path, max_count: int) -> tuple[list[FigureData], bool]:
    import docx

    document = docx.Document(str(path))
    figures: list[FigureData] = []
    for rel in document.part.rels.values():
        if len(figures) >= max_count:
            return figures, True  # truncated: another candidate remains
        if rel.is_external:
            continue
        try:
            content_type = rel.target_part.content_type
            blob: bytes = rel.target_part.blob
        except Exception:
            continue
        if not content_type or not content_type.startswith("image/"):
            continue
        ext = content_type.split("/")[-1].lower()
        figures.append(_make_figure_data(blob, ext, mime=content_type))
    return figures, False


def _make_figure_data(data: bytes, ext: str, mime: str | None = None) -> FigureData:
    """Build a ``FigureData``, deriving mime and parsing PNG/JPEG dimensions."""
    ext = (ext or "bin").lower()
    if mime is None:
        mime = _ext_to_mime(ext)
    width = height = None
    if ext == "png" or mime == "image/png":
        wh = _png_size(data)
        if wh:
            width, height = wh
    elif ext in ("jpg", "jpeg", "jpe") or mime == "image/jpeg":
        wh = _jpeg_size(data)
        if wh:
            width, height = wh
    return FigureData(
        bytes_data=data, ext=ext, mime=mime, width=width, height=height
    )


def _ext_to_mime(ext: str) -> str:
    mapping = {
        "png": "image/png",
        "jpg": "image/jpeg",
        "jpeg": "image/jpeg",
        "jpe": "image/jpeg",
        "gif": "image/gif",
        "bmp": "image/bmp",
        "tif": "image/tiff",
        "tiff": "image/tiff",
        "webp": "image/webp",
        "svg": "image/svg+xml",
    }
    return mapping.get(ext, f"image/{ext}")


def _png_size(data: bytes) -> tuple[int, int] | None:
    """Parse width/height from a PNG's IHDR chunk."""
    if len(data) < 24 or not data.startswith(b"\x89PNG\r\n\x1a\n"):
        return None
    width = int.from_bytes(data[16:20], "big")
    height = int.from_bytes(data[20:24], "big")
    return width, height


def _jpeg_size(data: bytes) -> tuple[int, int] | None:
    """Parse width/height from a JPEG by scanning SOF markers."""
    if len(data) < 4 or not data.startswith(b"\xff\xd8"):
        return None
    i = 2
    while i + 8 < len(data):
        if data[i] != 0xFF:
            i += 1
            continue
        marker = data[i + 1]
        # SOF markers: C0-C3, C5-C7, C9-CB, CD-CF (exclude C4/C8/CC)
        if 0xC0 <= marker <= 0xCF and marker not in (0xC4, 0xC8, 0xCC):
            height = int.from_bytes(data[i + 5 : i + 7], "big")
            width = int.from_bytes(data[i + 7 : i + 9], "big")
            return width, height
        # Skip this marker's segment (length is 2 bytes BE at i+2)
        if i + 3 >= len(data):
            return None
        seg_len = int.from_bytes(data[i + 2 : i + 4], "big")
        i += 2 + seg_len
    return None


# ---------- full extraction (text + metadata + figures) ----------


def _extract_content(path: Path, llm_client=None) -> tuple[str | None, str]:
    """Return ``(content, extraction_method)`` for any supported file.

    ``content`` is ``None`` only when text extraction fails (corrupt /
    unsupported) -- the caller then skips the item. ``extraction_method`` is
    ``text_layer`` | ``vision_llm`` | ``none``.
    """
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        return _extract_pdf_content(path, llm_client)
    text = extract_text(path)
    if text is None:
        return None, "none"
    return text, "text_layer"


def _extract_full(path: Path, llm_client=None) -> InfoItemData | None:
    """Extract text + metadata + figures, packed into ``InfoItemData.extra``.

    Returns ``None`` if text extraction fails. Title falls back to the filename
    when metadata has no title. ``extraction_method`` records how the body text
    was obtained (text_layer / vision_llm / none).
    """
    content, extraction_method = _extract_content(path, llm_client)
    if content is None:
        return None
    metadata = extract_metadata(path)
    first_page = _first_page_text(path)
    affiliation = extract_author_affiliation(first_page)
    figures, truncated = extract_figures(path, settings.max_figures_per_item)
    if truncated:
        _logger.warning(
            "图表数达到上限 %d，已截断: %s", settings.max_figures_per_item, path
        )
    title = metadata.get("title") or path.name
    mtime = datetime.fromtimestamp(path.stat().st_mtime, tz=timezone.utc)
    return InfoItemData(
        external_id=str(path.resolve()),
        title=title,
        url=str(path),
        content=content,
        published_at=mtime,
        extra={
            "author": metadata.get("author"),
            "author_affiliation": affiliation,
            "article_published_at": metadata.get("published_at"),
            "page_count": metadata.get("page_count"),
            "figures": figures,
            "extraction_method": extraction_method,
        },
    )


# ---------- adapter ----------


class LocalFolderAdapter(InfoSourceAdapter):
    type = "local_folder"

    def __init__(self, config: dict) -> None:
        super().__init__(config)
        self.folder_path: Path = Path(config["folder_path"])
        self.patterns: list[str] = config.get("patterns") or [
            "*.txt",
            "*.md",
            "*.pdf",
            "*.docx",
            "*.html",
        ]
        self.recursive: bool = bool(config.get("recursive", True))
        self.max_items: int = int(config.get("max_items") or 100000)

    @staticmethod
    def required_config_keys() -> list[str]:
        return ["folder_path"]

    def _iter_files(self):
        if not self.folder_path.exists():
            return
        glob = self.folder_path.rglob if self.recursive else self.folder_path.glob
        seen: set[str] = set()
        for pattern in self.patterns:
            for f in glob(pattern):
                if f.is_file():
                    key = str(f)
                    if key not in seen:
                        seen.add(key)
                        yield f

    def check_status(self) -> SourceStatus:
        if not self.folder_path.exists():
            return SourceStatus(ok=False, message=f"文件夹不存在: {self.folder_path}")
        count = sum(1 for _ in self._iter_files())
        return SourceStatus(ok=True, message=f"共 {count} 个匹配文件", item_count=count)

    def fetch_new_items(
        self,
        since: datetime | None = None,
        known_ids: set[str] | None = None,
    ) -> list[InfoItemData]:
        known = known_ids or set()
        # DB 读回的 last_sync_at 可能是 naive datetime，归一为 aware UTC 再与 mtime 比较。
        if since and since.tzinfo is None:
            since = since.replace(tzinfo=timezone.utc)
        items: list[InfoItemData] = []
        # 按路径排序，保证同步顺序确定、可解释（不依赖文件系统遍历顺序）。
        for f in sorted(self._iter_files()):
            ext_id = str(f.resolve())
            mtime = datetime.fromtimestamp(f.stat().st_mtime, tz=timezone.utc)
            # 增量 + 回补：已索引且未变更的文件跳过，不重读内容。
            if since and ext_id in known and mtime <= since:
                continue
            item = _extract_full(f)
            if item is None:
                continue
            items.append(item)
            if len(items) >= self.max_items:
                break
        return items

    def reextract(self, external_id: str) -> InfoItemData | None:
        """Re-extract a single file by its path (external_id)."""
        p = Path(external_id)
        if not p.exists():
            return None
        return _extract_full(p)
